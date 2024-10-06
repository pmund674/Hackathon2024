import tkinter as tk
from tkinter import messagebox, simpledialog
from tkinter import ttk
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import smtplib
import schedule
import time
import threading
from datetime import datetime, timedelta
import pandas as pd
from docx import Document
import winsound
from tkcalendar import Calendar
import os
import csv

# Assuming users are stored in a CSV file (username,password format)
USER_FILE = 'users.csv'

# Function to load users from a CSV file
def load_users():
    if not os.path.exists(USER_FILE):
        return {}
    with open(USER_FILE, mode='r', newline='') as f:
        reader = csv.reader(f)
        return {rows[0]: rows[1] for rows in reader}

# Function to save a new user to the CSV file
def save_user(username, password):
    with open(USER_FILE, mode='a', newline='') as f:
        writer = csv.writer(f)
        writer.writerow([username, password])

class Event:
    def __init__(self, start_hour, end_hour, name, email, description="", recurrence=1):
        self.start_hour = start_hour
        self.end_hour = end_hour
        self.name = name
        self.email = email
        self.description = description
        self.recurrence = recurrence
        self.date_time = None

class ScheduleBuilder:
    def __init__(self):
        self.schedule = {}

    def block_time(self, year, month, day, start_hour, end_hour, name, email, description="", recurrence=1):
        date = (year, month, day)
        event = Event(start_hour, end_hour, name, email, description, recurrence)
        event.date_time = datetime(year, month, day, start_hour)

        if date not in self.schedule:
            self.schedule[date] = []
        self.schedule[date].append(event)

        # Schedule email reminders
        self.schedule_email_reminders(event)
        
        # Handle recurrence
        for i in range(1, recurrence):
            self.block_time(year, month, day + i, start_hour, end_hour, name, email, description)
        return "Time blocked successfully!"

    def schedule_email_reminders(self, event):
        """Schedule emails to be sent 1 day and 10 minutes before the event."""
        event_time = event.date_time
        day_before = event_time - timedelta(days=1)
        ten_minutes_before = event_time - timedelta(minutes=10)

        # Scheduling day-before reminder
        schedule.every().day.at(day_before.strftime("%H:%M")).do(self.send_email_reminder, event, "Reminder: Your event is tomorrow!")

        # Scheduling 10-minutes-before reminder
        schedule.every().day.at(ten_minutes_before.strftime("%H:%M")).do(self.send_email_reminder, event, "Reminder: Your event is in 10 minutes!")

    def send_email_reminder(self, event, message):
        """Sends an email reminder."""
        try:
            self.play_notification_sound()  # Play notification sound before sending email
            sender_email = "sreeragvaddel@example.com"  # Replace with your email
            sender_password = "yourpassword"  # Replace with your email password

            # Create the email content
            msg = MIMEMultipart()
            msg['From'] = sender_email
            msg['To'] = event.email
            msg['Subject'] = f"Reminder: {event.name} event"

            body = f"{message}\n\nEvent: {event.name}\nTime: {event.start_hour}:00 - {event.end_hour}:00\nDescription: {event.description}"
            msg.attach(MIMEText(body, 'plain'))

            # Setup the SMTP server
            server = smtplib.SMTP('smtp.gmail.com', 587)  # Change this if not using Gmail
            server.starttls()
            server.login(sender_email, sender_password)
            text = msg.as_string()
            server.sendmail(sender_email, event.email, text)
            server.quit()

            print(f"Reminder email sent to {event.email}")
        except Exception as e:
            print(f"Failed to send email: {e}")

    def play_notification_sound(self):
        duration = 1000  # milliseconds
        frequency = 440  # Hz
        winsound.Beep(frequency, duration)

    def view_schedule(self, year, month, day):
        date = (year, month, day)
        if date in self.schedule:
            events = ["{}:00 - {}:00: {} (Email: {})".format(event.start_hour, event.end_hour, event.name, event.email) for event in self.schedule[date]]
            return "\n".join(events)
        else:
            return "No events scheduled for {}/{}/{}".format(month, day, year)

    def delete_event(self, year, month, day, start_hour, name):
        date = (year, month, day)
        if date in self.schedule:
            for event in self.schedule[date]:
                if event.start_hour == start_hour and event.name == name:
                    self.schedule[date].remove(event)
                    return "Event deleted successfully!"
            return "Event not found!"
        else:
            return "No events found for {}/{}/{}".format(month, day, year)

    def export_schedule(self, format="excel"):
        if format == "excel":
            data = []
            for date, events in self.schedule.items():
                for event in events:
                    data.append([date[0], date[1], date[2], event.start_hour, event.end_hour, event.name, event.email, event.description])
            df = pd.DataFrame(data, columns=["Year", "Month", "Day", "Start Hour", "End Hour", "Name", "Email", "Description"])
            df.to_excel("schedule.xlsx", index=False)
            return "Schedule exported to schedule.xlsx successfully!"
        
        elif format == "word":
            doc = Document()
            doc.add_heading('Schedule', level=1)

            for date, events in self.schedule.items():
                doc.add_heading(f"Date: {date[1]}/{date[2]}/{date[0]}", level=2)
                for event in events:
                    doc.add_paragraph(f"{event.start_hour}:00 - {event.end_hour}:00: {event.name} (Email: {event.email})")
                    if event.description:
                        doc.add_paragraph(f"Description: {event.description}")

            doc.save("schedule.docx")
            return "Schedule exported to schedule.docx successfully!"


class App:
    def __init__(self, root):
        self.builder = ScheduleBuilder()
        self.root = root
        self.root.title("Schedule Manager")

        self.users = load_users()
        self.current_user = None

        # Main Title
        self.label = tk.Label(root, text="Schedule Manager", font=('Helvetica', 16))
        self.label.pack(pady=10)

        # Create Account Button
        self.create_account_button = tk.Button(root, text="Create Account", command=self.create_account_gui)
        self.create_account_button.pack(pady=5)

        # Login Button
        self.login_button = tk.Button(root, text="Login", command=self.login_gui)
        self.login_button.pack(pady=5)

        # Block Time Button
        self.block_button = tk.Button(root, text="Block Time", command=self.block_time_gui, state='disabled')
        self.block_button.pack(pady=5)

        # View Schedule Button
        self.view_button = tk.Button(root, text="View Schedule", command=self.view_schedule_gui, state='disabled')
        self.view_button.pack(pady=5)

        # Delete Event Button
        self.delete_button = tk.Button(root, text="Delete Event", command=self.delete_event_gui, state='disabled')
        self.delete_button.pack(pady=5)

        # Export Schedule Button
        self.export_button = tk.Button(root, text="Export Schedule", command=self.export_schedule_gui, state='disabled')
        self.export_button.pack(pady=5)

    def create_account_gui(self):
        account_window = tk.Toplevel(self.root)
        account_window.title("Create Account")

        tk.Label(account_window, text="Username:").grid(row=0, column=0)
        username_entry = tk.Entry(account_window)
        username_entry.grid(row=0, column=1)

        tk.Label(account_window, text="Password:").grid(row=1, column=0)
        password_entry = tk.Entry(account_window, show="*")
        password_entry.grid(row=1, column=1)

        def create_account_action():
            username = username_entry.get()
            password = password_entry.get()
            if username in self.users:
                messagebox.showerror("Error", "Username already exists.")
            else:
                save_user(username, password)
                self.users[username] = password  # Update the in-memory user list
                messagebox.showinfo("Success", "Account created successfully!")
                account_window.destroy()

        create_account_button = tk.Button(account_window, text="Create Account", command=create_account_action)
        create_account_button.grid(row=2, columnspan=2, pady=10)

    def login_gui(self):
        login_window = tk.Toplevel(self.root)
        login_window.title("Login")

        tk.Label(login_window, text="Username:").grid(row=0, column=0)
        username_entry = tk.Entry(login_window)
        username_entry.grid(row=0, column=1)

        tk.Label(login_window, text="Password:").grid(row=1, column=0)
        password_entry = tk.Entry(login_window, show="*")
        password_entry.grid(row=1, column=1)

        def login_action():
            username = username_entry.get()
            password = password_entry.get()
            if username in self.users and self.users[username] == password:
                self.current_user = username
                messagebox.showinfo("Success", "Login successful!")
                login_window.destroy()
                self.block_button.config(state='normal')
                self.view_button.config(state='normal')
                self.delete_button.config(state='normal')
                self.export_button.config(state='normal')
            else:
                messagebox.showerror("Error", "Invalid username or password.")

        login_button = tk.Button(login_window, text="Login", command=login_action)
        login_button.grid(row=2, columnspan=2, pady=10)

    def block_time_gui(self):
        block_window = tk.Toplevel(self.root)
        block_window.title("Block Time")

        tk.Label(block_window, text="Select Date:").grid(row=0, column=0)
        cal = Calendar(block_window, selectmode='day')
        cal.grid(row=0, column=1)

        tk.Label(block_window, text="Start Hour:").grid(row=1, column=0)
        start_hour_entry = tk.Entry(block_window)
        start_hour_entry.grid(row=1, column=1)

        tk.Label(block_window, text="End Hour:").grid(row=2, column=0)
        end_hour_entry = tk.Entry(block_window)
        end_hour_entry.grid(row=2, column=1)

        tk.Label(block_window, text="Event Name:").grid(row=3, column=0)
        event_name_entry = tk.Entry(block_window)
        event_name_entry.grid(row=3, column=1)

        tk.Label(block_window, text="Email:").grid(row=4, column=0)
        email_entry = tk.Entry(block_window)
        email_entry.grid(row=4, column=1)

        tk.Label(block_window, text="Description:").grid(row=5, column=0)
        description_entry = tk.Entry(block_window)
        description_entry.grid(row=5, column=1)

        tk.Label(block_window, text="Recurrence (days):").grid(row=6, column=0)
        recurrence_entry = tk.Entry(block_window)
        recurrence_entry.grid(row=6, column=1)

        def block_time_action():
            date = cal.selection_get()
            year, month, day = date.year, date.month, date.day
            start_hour = int(start_hour_entry.get())
            end_hour = int(end_hour_entry.get())
            name = event_name_entry.get()
            email = email_entry.get()
            description = description_entry.get()
            recurrence = int(recurrence_entry.get()) if recurrence_entry.get() else 1

            result = self.builder.block_time(year, month, day, start_hour, end_hour, name, email, description, recurrence)
            messagebox.showinfo("Info", result)

        block_button = tk.Button(block_window, text="Block Time", command=block_time_action)
        block_button.grid(row=7, columnspan=2, pady=10)

    def view_schedule_gui(self):
        view_window = tk.Toplevel(self.root)
        view_window.title("View Schedule")

        tk.Label(view_window, text="Select Date:").grid(row=0, column=0)
        cal = Calendar(view_window, selectmode='day')
        cal.grid(row=0, column=1)

        def view_schedule_action():
            date = cal.selection_get()
            year, month, day = date.year, date.month, date.day
            schedule_text = self.builder.view_schedule(year, month, day)
            messagebox.showinfo("Schedule", schedule_text)

        view_button = tk.Button(view_window, text="View Schedule", command=view_schedule_action)
        view_button.grid(row=1, columnspan=2, pady=10)

    def delete_event_gui(self):
        delete_window = tk.Toplevel(self.root)
        delete_window.title("Delete Event")

        tk.Label(delete_window, text="Select Date:").grid(row=0, column=0)
        cal = Calendar(delete_window, selectmode='day')
        cal.grid(row=0, column=1)

        tk.Label(delete_window, text="Start Hour:").grid(row=1, column=0)
        start_hour_entry = tk.Entry(delete_window)
        start_hour_entry.grid(row=1, column=1)

        tk.Label(delete_window, text="Event Name:").grid(row=2, column=0)
        event_name_entry = tk.Entry(delete_window)
        event_name_entry.grid(row=2, column=1)

        def delete_event_action():
            date = cal.selection_get()
            year, month, day = date.year, date.month, date.day
            start_hour = int(start_hour_entry.get())
            name = event_name_entry.get()

            result = self.builder.delete_event(year, month, day, start_hour, name)
            messagebox.showinfo("Info", result)

        delete_button = tk.Button(delete_window, text="Delete Event", command=delete_event_action)
        delete_button.grid(row=3, columnspan=2, pady=10)

    def export_schedule_gui(self):
        export_window = tk.Toplevel(self.root)
        export_window.title("Export Schedule")

        tk.Label(export_window, text="Select Format:").grid(row=0, column=0)
        format_var = tk.StringVar(value="excel")
        tk.Radiobutton(export_window, text="Excel", variable=format_var, value="excel").grid(row=0, column=1)
        tk.Radiobutton(export_window, text="Word", variable=format_var, value="word").grid(row=1, column=1)

        def export_schedule_action():
            format_selected = format_var.get()
            result = self.builder.export_schedule(format=format_selected)
            messagebox.showinfo("Info", result)

        export_button = tk.Button(export_window, text="Export", command=export_schedule_action)
        export_button.grid(row=2, columnspan=2, pady=10)

    def run_scheduler(self):
        while True:
            schedule.run_pending()
            time.sleep(1)

if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)

    # Start the scheduler thread
    scheduler_thread = threading.Thread(target=app.run_scheduler, daemon=True)
    scheduler_thread.start()

    root.mainloop()
# you need these pip installed in your terminal to  run this code
#     pip install tk
#     pip install tkcalendar pandas python-docx
#     pip install pandas openpyxl python-docx
#     pip install schedule